<<<<<<< HEAD


import re 
from docx import Document 
from docx .shared import Pt ,Inches 
from docx .enum .text import WD_ALIGN_PARAGRAPH 


def parse_quiz_text (quiz_text ):
    
    data ={"title":"Quiz","subtitle":"","source":"General knowledge","questions":[]}

    title_match =re .search (r"TITLE:\s*(.+)",quiz_text )
    subtitle_match =re .search (r"SUBTITLE:\s*(.+)",quiz_text )
    source_match =re .search (r"SOURCE:\s*(.+)",quiz_text )

    if title_match :
        data ["title"]=title_match .group (1 ).strip ()
    if subtitle_match :
        data ["subtitle"]=subtitle_match .group (1 ).strip ()
    if source_match :
        data ["source"]=source_match .group (1 ).strip ()


    blocks =re .split (r"\n(?=Q\d+:)",quiz_text )

    for block in blocks :
        q_match =re .search (r"Q(\d+):\s*(.+)",block )
        if not q_match :
            continue 

        q_num =q_match .group (1 )
        q_text =q_match .group (2 ).strip ()

        type_match =re .search (r"TYPE:\s*(.+?)(?=\n|$)",block )


        answer_match =re .search (r"ANSWER:\s*(.+?)(?=EXPLANATION:|$)",block ,re .DOTALL )


        explanation_match =re .search (r"EXPLANATION:\s*(.+?)(?=Q\d+:|$)",block ,re .DOTALL )

        options =re .findall (r"^([A-Z]\))\s*(.+)$",block ,re .MULTILINE )


        answer_text =answer_match .group (1 ).strip ()if answer_match else ""
        explanation_text =explanation_match .group (1 ).strip ()if explanation_match else ""


        answer_text =answer_text .rstrip (')**')
        explanation_text =explanation_text .rstrip (')**')

        data ["questions"].append ({
        "number":q_num ,
        "question":q_text ,
        "type":type_match .group (1 ).strip ()if type_match else "SHORT_ANSWER",
        "options":options ,
        "answer":answer_text ,
        "explanation":explanation_text ,
        })

    return data 


def create_quiz_document (quiz_text ,output_path ="quiz.docx",include_answers =False ):
    
    data =parse_quiz_text (quiz_text )
    doc =Document ()


    title =doc .add_heading (data ["title"],level =0 )
    title .alignment =WD_ALIGN_PARAGRAPH .CENTER 

    if data ["subtitle"]:
        subtitle =doc .add_paragraph (data ["subtitle"])
        subtitle .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        subtitle .runs [0 ].italic =True 
        subtitle .runs [0 ].font .size =Pt (12 )


    if data .get ("source"):
        source_para =doc .add_paragraph (f"Source: {data['source']}")
        source_para .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        source_para .runs [0 ].italic =True 
        source_para .runs [0 ].font .size =Pt (9 )

    doc .add_paragraph ()


    instructions =doc .add_paragraph ("Answer each question to the best of your ability.")
    instructions .runs [0 ].italic =True 

    doc .add_paragraph ()


    for q in data ["questions"]:
        q_para =doc .add_paragraph ()
        q_run =q_para .add_run (f"{q['number']}. {q['question']}")
        q_run .bold =True 
        q_run .font .size =Pt (12 )

        if q ["options"]:
            for label ,text in q ["options"]:
                opt_para =doc .add_paragraph (f"{label} {text}")
                opt_para .paragraph_format .left_indent =Inches (0.3 )
        elif q ["type"]=="TRUE_FALSE":
            doc .add_paragraph ("☐ True     ☐ False").paragraph_format .left_indent =Inches (0.3 )
        elif q ["type"]=="FILL_BLANK":
            pass 
        elif q ["type"]=="SHORT_ANSWER":
            for _ in range (2 ):
                line =doc .add_paragraph ("_"*60 )
                line .paragraph_format .left_indent =Inches (0.3 )

        doc .add_paragraph ()


    if include_answers :
        doc .add_page_break ()
        answer_heading =doc .add_heading ("Answer Key",level =1 )
        answer_heading .alignment =WD_ALIGN_PARAGRAPH .CENTER 

        for q in data ["questions"]:
            ans_para =doc .add_paragraph ()
            ans_run =ans_para .add_run (f"{q['number']}. Answer: {q['answer']}")
            ans_run .bold =True 

            if q ["explanation"]:
                exp_para =doc .add_paragraph (q ["explanation"])
                exp_para .paragraph_format .left_indent =Inches (0.3 )
                exp_para .runs [0 ].italic =True 

            doc .add_paragraph ()

    doc .save (output_path )
    return output_path 


def create_quiz_and_answer_key (quiz_text ,quiz_path ="quiz.docx",answer_key_path ="answer_key.docx"):
    
    create_quiz_document (quiz_text ,output_path =quiz_path ,include_answers =False )


    data =parse_quiz_text (quiz_text )
    doc =Document ()

    heading =doc .add_heading (f"{data['title']} - Answer Key",level =0 )
    heading .alignment =WD_ALIGN_PARAGRAPH .CENTER 

    if data .get ("source"):
        source_para =doc .add_paragraph (f"Source: {data['source']}")
        source_para .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        source_para .runs [0 ].italic =True 
        source_para .runs [0 ].font .size =Pt (9 )
        doc .add_paragraph ()

    for q in data ["questions"]:
        ans_para =doc .add_paragraph ()
        ans_run =ans_para .add_run (f"{q['number']}. {q['question']}")
        ans_run .bold =True 

        answer_para =doc .add_paragraph (f"Correct answer: {q['answer']}")
        answer_para .paragraph_format .left_indent =Inches (0.3 )

        if q ["explanation"]:
            exp_para =doc .add_paragraph (q ["explanation"])
            exp_para .paragraph_format .left_indent =Inches (0.3 )
            exp_para .runs [0 ].italic =True 

        doc .add_paragraph ()

    doc .save (answer_key_path )
=======


import re 
from docx import Document 
from docx .shared import Pt ,Inches 
from docx .enum .text import WD_ALIGN_PARAGRAPH 


def parse_quiz_text (quiz_text ):
    
    data ={"title":"Quiz","subtitle":"","source":"General knowledge","questions":[]}

    title_match =re .search (r"TITLE:\s*(.+)",quiz_text )
    subtitle_match =re .search (r"SUBTITLE:\s*(.+)",quiz_text )
    source_match =re .search (r"SOURCE:\s*(.+)",quiz_text )

    if title_match :
        data ["title"]=title_match .group (1 ).strip ()
    if subtitle_match :
        data ["subtitle"]=subtitle_match .group (1 ).strip ()
    if source_match :
        data ["source"]=source_match .group (1 ).strip ()


    blocks =re .split (r"\n(?=Q\d+:)",quiz_text )

    for block in blocks :
        q_match =re .search (r"Q(\d+):\s*(.+)",block )
        if not q_match :
            continue 

        q_num =q_match .group (1 )
        q_text =q_match .group (2 ).strip ()

        type_match =re .search (r"TYPE:\s*(.+?)(?=\n|$)",block )


        answer_match =re .search (r"ANSWER:\s*(.+?)(?=EXPLANATION:|$)",block ,re .DOTALL )


        explanation_match =re .search (r"EXPLANATION:\s*(.+?)(?=Q\d+:|$)",block ,re .DOTALL )

        options =re .findall (r"^([A-Z]\))\s*(.+)$",block ,re .MULTILINE )


        answer_text =answer_match .group (1 ).strip ()if answer_match else ""
        explanation_text =explanation_match .group (1 ).strip ()if explanation_match else ""


        answer_text =answer_text .rstrip (')**')
        explanation_text =explanation_text .rstrip (')**')

        data ["questions"].append ({
        "number":q_num ,
        "question":q_text ,
        "type":type_match .group (1 ).strip ()if type_match else "SHORT_ANSWER",
        "options":options ,
        "answer":answer_text ,
        "explanation":explanation_text ,
        })

    return data 


def create_quiz_document (quiz_text ,output_path ="quiz.docx",include_answers =False ):
    
    data =parse_quiz_text (quiz_text )
    doc =Document ()


    title =doc .add_heading (data ["title"],level =0 )
    title .alignment =WD_ALIGN_PARAGRAPH .CENTER 

    if data ["subtitle"]:
        subtitle =doc .add_paragraph (data ["subtitle"])
        subtitle .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        subtitle .runs [0 ].italic =True 
        subtitle .runs [0 ].font .size =Pt (12 )


    if data .get ("source"):
        source_para =doc .add_paragraph (f"Source: {data['source']}")
        source_para .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        source_para .runs [0 ].italic =True 
        source_para .runs [0 ].font .size =Pt (9 )

    doc .add_paragraph ()


    instructions =doc .add_paragraph ("Answer each question to the best of your ability.")
    instructions .runs [0 ].italic =True 

    doc .add_paragraph ()


    for q in data ["questions"]:
        q_para =doc .add_paragraph ()
        q_run =q_para .add_run (f"{q['number']}. {q['question']}")
        q_run .bold =True 
        q_run .font .size =Pt (12 )

        if q ["options"]:
            for label ,text in q ["options"]:
                opt_para =doc .add_paragraph (f"{label} {text}")
                opt_para .paragraph_format .left_indent =Inches (0.3 )
        elif q ["type"]=="TRUE_FALSE":
            doc .add_paragraph ("☐ True     ☐ False").paragraph_format .left_indent =Inches (0.3 )
        elif q ["type"]=="FILL_BLANK":
            pass 
        elif q ["type"]=="SHORT_ANSWER":
            for _ in range (2 ):
                line =doc .add_paragraph ("_"*60 )
                line .paragraph_format .left_indent =Inches (0.3 )

        doc .add_paragraph ()


    if include_answers :
        doc .add_page_break ()
        answer_heading =doc .add_heading ("Answer Key",level =1 )
        answer_heading .alignment =WD_ALIGN_PARAGRAPH .CENTER 

        for q in data ["questions"]:
            ans_para =doc .add_paragraph ()
            ans_run =ans_para .add_run (f"{q['number']}. Answer: {q['answer']}")
            ans_run .bold =True 

            if q ["explanation"]:
                exp_para =doc .add_paragraph (q ["explanation"])
                exp_para .paragraph_format .left_indent =Inches (0.3 )
                exp_para .runs [0 ].italic =True 

            doc .add_paragraph ()

    doc .save (output_path )
    return output_path 


def create_quiz_and_answer_key (quiz_text ,quiz_path ="quiz.docx",answer_key_path ="answer_key.docx"):
    
    create_quiz_document (quiz_text ,output_path =quiz_path ,include_answers =False )


    data =parse_quiz_text (quiz_text )
    doc =Document ()

    heading =doc .add_heading (f"{data['title']} - Answer Key",level =0 )
    heading .alignment =WD_ALIGN_PARAGRAPH .CENTER 

    if data .get ("source"):
        source_para =doc .add_paragraph (f"Source: {data['source']}")
        source_para .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        source_para .runs [0 ].italic =True 
        source_para .runs [0 ].font .size =Pt (9 )
        doc .add_paragraph ()

    for q in data ["questions"]:
        ans_para =doc .add_paragraph ()
        ans_run =ans_para .add_run (f"{q['number']}. {q['question']}")
        ans_run .bold =True 

        answer_para =doc .add_paragraph (f"Correct answer: {q['answer']}")
        answer_para .paragraph_format .left_indent =Inches (0.3 )

        if q ["explanation"]:
            exp_para =doc .add_paragraph (q ["explanation"])
            exp_para .paragraph_format .left_indent =Inches (0.3 )
            exp_para .runs [0 ].italic =True 

        doc .add_paragraph ()

    doc .save (answer_key_path )
>>>>>>> 6696ff70c425dd6f93af6c93d97bcaa324f38300
    return quiz_path ,answer_key_path 